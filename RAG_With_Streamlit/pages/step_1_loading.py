# modified_document_loader.py
import os
import json
import csv
from docx import Document
from PyPDF2 import PdfReader
import pandas as pd

def load_documents_from_folder(folder_path):
    """
    Load all supported documents (TXT, DOCX, PDF, CSV, JSON, XLSX, MD)
    from a given folder.
    Args: folder_path (str): Path to the folder containing documents.
    Returns: list: List of dictionaries with keys:
              'content', 'source', 'length', and 'file_type'
    """
    print("=" * 60)
    print("STEP 1: Loading documents from a folder")
    print("=" * 60)

    supported_exts = {".txt", ".docx", ".pdf", ".csv", ".json", ".xlsx", ".md"}
    documents = []

    if not os.path.isdir(folder_path):
        print(f"Error: {folder_path} is not a valid directory.")
        return documents

    file_paths = [
        os.path.join(folder_path, f)
        for f in os.listdir(folder_path)
        if os.path.splitext(f)[1].lower() in supported_exts
    ]

    if not file_paths:
        print("No supported files found in the folder.")
        return documents

    for file_path in file_paths:
        ext = os.path.splitext(file_path)[1].lower()
        content = ""

        try:
            # ----- TXT -----
            if ext == ".txt":
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()

            # ----- DOCX -----
            elif ext == ".docx":
                doc = Document(file_path)
                content = "\n".join([para.text for para in doc.paragraphs])

            # ----- PDF -----
            elif ext == ".pdf":
                reader = PdfReader(file_path)
                pages = [page.extract_text() or "" for page in reader.pages]
                content = "\n".join(pages)

            # ----- CSV -----
            elif ext == ".csv":
                rows = []
                with open(file_path, "r", encoding="utf-8") as f:
                    reader = csv.reader(f)
                    header = next(reader, None)
                    for row in reader:
                        if header:
                            rows.append(dict(zip(header, row)))
                        else:
                            rows.append(row)
                content = json.dumps(rows, indent=2, ensure_ascii=False)

            # ----- JSON -----
            elif ext == ".json":
                with open(file_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                content = json.dumps(data, indent=2, ensure_ascii=False)

            # ----- XLSX -----
            elif ext == ".xlsx":
                df = pd.read_excel(file_path)
                content = df.to_json(orient="records", indent=2, force_ascii=False)

            # ----- Markdown -----
            elif ext == ".md":
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()

            else:
                print(f"Skipping unsupported file type: {file_path}")
                continue

            # Append document info
            documents.append({
                "content": content,
                "source": file_path,
                "length": len(content),
                "file_type": ext.replace('.', '')
            })

            print(f"Loaded: {file_path}")
            print(f"  - Type: {ext}")
            print(f"  - Characters: {len(content)}")
            print(f"  - Words: {len(content.split())}")

        except Exception as e:
            print(f"Error loading {file_path}: {e}")

    print(f"\nTotal documents loaded: {len(documents)}")
    return documents



def load_documents_from_streamlit_files(uploaded_files):
    """
    Modified version of the above function that works directly with Streamlit uploaded files without needing to pass the path of folder
    Load all supported documents directly from Streamlit uploaded files
    """
    print("=" * 60)
    print("STEP 1: Loading documents from Streamlit uploaded files")
    print("=" * 60)

    documents = []

    if not uploaded_files:
        print("No files provided.")
        return documents

    for uploaded_file in uploaded_files:
        file_name = uploaded_file.name
        ext = os.path.splitext(file_name)[1].lower()
        content = ""

        try:
            # ----- TXT -----
            if ext == ".txt":
                content = uploaded_file.getvalue().decode("utf-8")

            # ----- DOCX -----
            elif ext == ".docx":
                # Reset file pointer to beginning
                uploaded_file.seek(0)
                doc = Document(uploaded_file)
                content = "\n".join([para.text for para in doc.paragraphs])

            # ----- PDF -----
            elif ext == ".pdf":
                # Reset file pointer to beginning
                uploaded_file.seek(0)
                reader = PdfReader(uploaded_file)
                pages = [page.extract_text() or "" for page in reader.pages]
                content = "\n".join(pages)

            # ----- CSV -----
            elif ext == ".csv":
                # Reset file pointer to beginning
                uploaded_file.seek(0)
                text_content = uploaded_file.getvalue().decode("utf-8")
                reader = csv.reader(text_content.splitlines())
                rows = []
                header = next(reader, None)
                for row in reader:
                    if header:
                        rows.append(dict(zip(header, row)))
                    else:
                        rows.append(row)
                content = json.dumps(rows, indent=2, ensure_ascii=False)

            # ----- JSON -----
            elif ext == ".json":
                # Reset file pointer to beginning
                uploaded_file.seek(0)
                data = json.load(uploaded_file)
                content = json.dumps(data, indent=2, ensure_ascii=False)

            # ----- XLSX -----
            elif ext == ".xlsx":
                # Reset file pointer to beginning
                uploaded_file.seek(0)
                df = pd.read_excel(uploaded_file)
                content = df.to_json(orient="records", indent=2, force_ascii=False)

            # ----- Markdown -----
            elif ext == ".md":
                content = uploaded_file.getvalue().decode("utf-8")

            else:
                print(f"Skipping unsupported file type: {file_name}")
                continue

            # Append document info
            documents.append({
                "content": content,
                "source": file_name,
                "length": len(content),
                "file_type": ext.replace('.', '')
            })

            print(f"Loaded: {file_name}")
            print(f"  - Type: {ext}")
            print(f"  - Characters: {len(content)}")
            print(f"  - Words: {len(content.split())}")

        except Exception as e:
            print(f"Error loading {file_name}: {e}")

    print(f"\nTotal documents loaded: {len(documents)}")
    return documents




